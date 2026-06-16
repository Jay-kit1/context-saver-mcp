from __future__ import annotations

import csv
import json
import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from .file_types import detect_file_type, is_supported_file


def _read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="replace")


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for index, page in enumerate(reader.pages, start=1):
        parts.append(f"[Page {index}]\n{page.extract_text() or ''}")
    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    table_parts = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        table_parts.append(f"[Table {table_index}]\n" + "\n".join(rows))
    return "\n\n".join(paragraphs + table_parts)


def _extract_csv(path: Path) -> str:
    text = _read_text(path)
    rows = list(csv.reader(text.splitlines()))
    if not rows:
        return ""
    sample = rows[:21]
    return "\n".join(",".join(row) for row in sample)


def _extract_json(path: Path) -> str:
    text = _read_text(path)
    if len(text) < 8000:
        return text
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        return text[:8000]
    if isinstance(data, dict):
        keys = list(data.keys())[:80]
        return "JSON object keys:\n" + "\n".join(f"- {key}: {type(data[key]).__name__}" for key in keys)
    if isinstance(data, list):
        return f"JSON array length: {len(data)}\nSample:\n{json.dumps(data[:5], ensure_ascii=False, indent=2)}"
    return str(data)


def _extract_log(path: Path) -> str:
    lines = _read_text(path).splitlines()
    interesting = []
    pattern = re.compile(r"error|warning|traceback|exception", re.I)
    for index, line in enumerate(lines):
        if pattern.search(line):
            start = max(0, index - 3)
            end = min(len(lines), index + 6)
            interesting.append(f"[lines {start + 1}-{end}]\n" + "\n".join(lines[start:end]))
    if interesting:
        return "\n\n".join(interesting[:80])
    return "\n".join(lines[:500])


def _extract_code(path: Path) -> str:
    text = _read_text(path)
    lines = text.splitlines()
    interesting = []
    patterns = (
        r"^\s*(import|from|require\(|class |def |function |const |let |var |export |interface |type )",
        r"(TODO|FIXME|ERROR|WARNING|Exception|Traceback|process\.env|DATABASE|TOKEN|SECRET|API_KEY)",
    )
    compiled = [re.compile(p) for p in patterns]
    for index, line in enumerate(lines):
        if any(regex.search(line) for regex in compiled):
            start = max(0, index - 2)
            end = min(len(lines), index + 4)
            interesting.append(f"[lines {start + 1}-{end}]\n" + "\n".join(lines[start:end]))
    if interesting and len(text) > 12000:
        return "\n\n".join(interesting[:120])
    return text


def extract_text_from_file(path: Path) -> str:
    path = Path(path)
    if not is_supported_file(path):
        raise ValueError(f"Unsupported or skipped file: {path}")
    file_type = detect_file_type(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".csv":
        return _extract_csv(path)
    if suffix == ".json":
        return _extract_json(path)
    if suffix == ".log":
        return _extract_log(path)
    if file_type == "code":
        return _extract_code(path)
    return _read_text(path)
